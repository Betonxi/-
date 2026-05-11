"""报告生成工具函数：格式辅助 + 图表生成"""
from __future__ import annotations
from pathlib import Path
import numpy as np
import matplotlib
matplotlib.use("Agg")
matplotlib.rcParams["font.sans-serif"] = ["Microsoft YaHei", "SimHei"]
matplotlib.rcParams["axes.unicode_minus"] = False
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml

FONT_SONG, FONT_HEI, FONT_EN = "宋体", "黑体", "Times New Roman"
PT_XIAOER, PT_SAN, PT_SI = Pt(18), Pt(16), Pt(14)
PT_XIAOSI, PT_WU = Pt(12), Pt(10.5)
COLOR_BLACK = RGBColor(0, 0, 0)

# ── 段落/字体 ──
def _set_run(run, font_cn=FONT_SONG, size=PT_XIAOSI, bold=False):
    run.font.name = FONT_EN
    run.font.size = size
    run.bold = bold
    run.font.color.rgb = COLOR_BLACK
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_cn)

def _fmt(para, indent=None, spacing=1.5, before=Pt(0), after=Pt(0),
         align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    f = para.paragraph_format
    f.alignment = align; f.space_before = before; f.space_after = after
    f.line_spacing_rule = WD_LINE_SPACING.MULTIPLE; f.line_spacing = spacing
    if indent: f.first_line_indent = indent

def _style(doc, name):
    return doc.styles[name]

def configure_document_styles(doc):
    s = _style(doc, "Heading 1")
    s.font.name = FONT_EN
    s.font.size = PT_SAN
    s.font.bold = True
    s.font.color.rgb = COLOR_BLACK
    s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.paragraph_format.space_before = Pt(12)
    s.paragraph_format.space_after = Pt(6)

    s = _style(doc, "Heading 2")
    s.font.name = FONT_EN
    s.font.size = PT_SI
    s.font.bold = True
    s.font.color.rgb = COLOR_BLACK
    s.paragraph_format.space_before = Pt(8)
    s.paragraph_format.space_after = Pt(4)

    s = _style(doc, "Heading 3")
    s.font.name = FONT_EN
    s.font.size = PT_XIAOSI
    s.font.bold = True
    s.font.color.rgb = COLOR_BLACK
    s.paragraph_format.space_before = Pt(6)
    s.paragraph_format.space_after = Pt(3)

    s = _style(doc, "Body Text")
    s.font.name = FONT_EN
    s.font.size = PT_XIAOSI
    s.font.color.rgb = COLOR_BLACK
    s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    s.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    s.paragraph_format.line_spacing = 1.5
    s.paragraph_format.space_after = Pt(2)
    s.paragraph_format.first_line_indent = Cm(0.74)

    s = _style(doc, "Caption")
    s.font.name = FONT_EN
    s.font.size = PT_WU
    s.font.color.rgb = COLOR_BLACK
    s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.paragraph_format.space_before = Pt(2)
    s.paragraph_format.space_after = Pt(6)

    s = _style(doc, "TOC Heading")
    s.font.name = FONT_EN
    s.font.size = PT_SAN
    s.font.bold = True
    s.font.color.rgb = COLOR_BLACK
    s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.paragraph_format.space_before = Pt(12)
    s.paragraph_format.space_after = Pt(6)

def clear_document_body(doc):
    body = doc._element.body
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)

def add_toc(doc):
    p = doc.add_paragraph(style="TOC Heading")
    r = p.add_run("目  录")
    _set_run(r, FONT_HEI, PT_SAN, True)

    p = doc.add_paragraph()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), 'TOC \\o "1-3" \\h \\z \\u')
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "右键更新目录"
    run.append(text)
    fld.append(run)
    p._p.append(fld)
    _fmt(p, indent=None, after=Pt(6), align=WD_ALIGN_PARAGRAPH.CENTER)
    return p

def enable_update_fields_on_open(doc):
    settings = doc.settings.element
    for child in settings:
        if child.tag.endswith("updateFields"):
            child.set(qn("w:val"), "true")
            return
    el = OxmlElement("w:updateFields")
    el.set(qn("w:val"), "true")
    settings.append(el)

def add_body(doc, text, bold=False, indent=Cm(0.74), size=PT_XIAOSI, font_cn=FONT_SONG,
             after=Pt(2), align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph(style="Body Text"); r = p.add_run(text)
    _set_run(r, font_cn=font_cn, size=size, bold=bold)
    _fmt(p, indent=indent, after=after, align=align); return p

def add_h1(doc, t):
    p=doc.add_paragraph(style="Heading 1"); r=p.add_run(t)
    _set_run(r, FONT_HEI, PT_SAN, True)
    return p

def add_h2(doc, t):
    p=doc.add_paragraph(style="Heading 2"); r=p.add_run(t)
    _set_run(r, FONT_HEI, PT_SI, True)
    return p

def add_h3(doc, t):
    p=doc.add_paragraph(style="Heading 3"); r=p.add_run(t)
    _set_run(r, FONT_HEI, PT_XIAOSI, True)
    return p

# ── 三线表 ──
def add_table3(doc, headers, rows, caption=None, cap_num=1):
    if caption:
        cp=doc.add_paragraph(style="Caption"); r=cp.add_run(f"表 {cap_num}  {caption}")
        _set_run(r, FONT_SONG, PT_WU, True)
        _fmt(cp, before=Pt(6), after=Pt(3), align=WD_ALIGN_PARAGRAPH.CENTER)
    nc=len(headers); tbl=doc.add_table(rows=1+len(rows), cols=nc)
    tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
    tbl._tbl.tblPr.append(parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="nil"/><w:left w:val="nil"/><w:bottom w:val="nil"/>'
        '<w:right w:val="nil"/><w:insideH w:val="nil"/><w:insideV w:val="nil"/>'
        '</w:tblBorders>'))
    def _cell_bdr(cell, top_sz=0, bot_sz=0):
        tc=cell._element.get_or_add_tcPr()
        t=f'w:val="single" w:sz="{top_sz}" w:space="0" w:color="000000"' if top_sz else 'w:val="nil"'
        b=f'w:val="single" w:sz="{bot_sz}" w:space="0" w:color="000000"' if bot_sz else 'w:val="nil"'
        tc.append(parse_xml(f'<w:tcBorders {nsdecls("w")}><w:top {t}/><w:bottom {b}/>'
                            '<w:left w:val="nil"/><w:right w:val="nil"/></w:tcBorders>'))
    for j,h in enumerate(headers):
        c=tbl.rows[0].cells[j]; c.text=""
        r=c.paragraphs[0].add_run(h); _set_run(r, FONT_HEI, PT_WU, True)
        c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
        _cell_bdr(c, 12, 8)
    for i,rd in enumerate(rows):
        last=(i==len(rows)-1)
        for j,v in enumerate(rd):
            c=tbl.rows[i+1].cells[j]; c.text=""
            r=c.paragraphs[0].add_run(str(v)); _set_run(r, FONT_SONG, PT_WU)
            c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
            _cell_bdr(c, 0, 12 if last else 0)
    return tbl

def add_fig(doc, path, caption, num, width=Cm(14)):
    pp=doc.add_paragraph(); pp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    pp.add_run().add_picture(str(path), width=width)
    cp=doc.add_paragraph(style="Caption"); r=cp.add_run(f"图 {num}  {caption}")
    _set_run(r, FONT_SONG, PT_WU)
    _fmt(cp, before=Pt(2), after=Pt(6), align=WD_ALIGN_PARAGRAPH.CENTER)

# ── 计数器 ──
_FIG=[0]; _TBL=[0]
def nf(): _FIG[0]+=1; return _FIG[0]
def nt(): _TBL[0]+=1; return _TBL[0]

# ── 图表生成 ──
def _box(ax, xy, w, h, text, fc="#e8f0fe", ec="#4285f4", fs=9):
    ax.add_patch(FancyBboxPatch(xy, w, h, boxstyle="round,pad=0.15",
                                fc=fc, ec=ec, lw=0.8, zorder=2))
    ax.text(xy[0]+w/2, xy[1]+h/2, text, ha="center", va="center", fontsize=fs, fontweight="bold", zorder=3)

def _arr(ax, s, e, c="#555"):
    ax.annotate("", xy=e, xytext=s, arrowprops=dict(arrowstyle="-|>", color=c, lw=1.2))

CHART_DIR = Path(__file__).resolve().parent / "_report_charts"
CHART_DIR.mkdir(exist_ok=True)

def gen_architecture():
    p=CHART_DIR/"fig_architecture.png"
    fig,ax=plt.subplots(figsize=(10,7),dpi=150); ax.set_xlim(0,10); ax.set_ylim(0,7.5); ax.axis("off")
    ax.set_title("FinChat 系统四层架构", fontsize=14, fontweight="bold", pad=10)
    for x,y,w,h,fc,lb in [(0.3,5.8,9.4,1.3,"#e8f4fd","前端交互层"),(0.3,4.2,9.4,1.3,"#e8fde8","API 服务层"),
                           (0.3,2.2,9.4,1.7,"#fef3e8","智能体编排层"),(0.3,0.3,9.4,1.6,"#f3e8fe","数据引擎层")]:
        ax.add_patch(FancyBboxPatch((x,y),w,h,boxstyle="round,pad=0.2",fc=fc,ec="#bbb",lw=0.8,alpha=0.6))
        ax.text(x+0.15,y+h-0.2,lb,fontsize=8,color="#666",fontstyle="italic",va="top")
    _box(ax,(1.2,6.0),2.2,0.8,"Web 工作台\nHTML+Tailwind",fc="#bbdefb")
    _box(ax,(4.0,6.0),2.2,0.8,"Plotly.js 图表",fc="#bbdefb")
    _box(ax,(6.8,6.0),2.2,0.8,"SSE 流式渲染",fc="#bbdefb")
    _box(ax,(2.0,4.4),2.5,0.8,"FastAPI\nRESTful+Stream",fc="#c8e6c9")
    _box(ax,(5.5,4.4),2.5,0.8,"Pydantic 校验",fc="#c8e6c9")
    _box(ax,(3.5,3.15),3.0,0.65,"Orchestrator 编排智能体",fc="#ffe0b2",ec="#e65100",fs=10)
    for x,y,t in [(0.6,2.35,"QueryAgent\n智能问数"),(3.0,2.35,"Assessment\n运营评估"),
                   (5.4,2.35,"RiskAgent\n风险洞察"),(7.8,2.35,"ReportAgent\n报告生成")]:
        _box(ax,(x,y),1.8,0.65,t,fc="#fff3e0",ec="#ef6c00",fs=8)
    _box(ax,(1.0,0.5),2.2,0.9,"SQLite 双数据库\n交叉验证",fc="#e1bee7")
    _box(ax,(4.0,0.5),2.2,0.9,"TF-IDF 知识库\n研报向量索引",fc="#e1bee7")
    _box(ax,(7.0,0.5),2.2,0.9,"共享黑板\nBlackboard",fc="#e1bee7")
    _arr(ax,(5,5.95),(5,5.25)); _arr(ax,(5,4.35),(5,3.85))
    fig.savefig(p,bbox_inches="tight",facecolor="white"); plt.close(fig); return p

def gen_agent_flow():
    p=CHART_DIR/"fig_agent_flow.png"
    fig,ax=plt.subplots(figsize=(12,4),dpi=150); ax.set_xlim(0,12); ax.set_ylim(0,4); ax.axis("off")
    ax.set_title("多智能体协同报告生成流程", fontsize=13, fontweight="bold", pad=10)
    for x,y,t,fc,ec in [(0.3,1.5,"用户请求","#e3f2fd","#1565c0"),(2.5,1.5,"Orchestrator\n意图路由","#fff8e1","#f57f17"),
                         (4.7,2.5,"QueryAgent","#e8f5e9","#2e7d32"),(4.7,0.5,"Assessment","#fce4ec","#c62828"),
                         (7.2,2.5,"RiskAgent","#f3e5f5","#6a1b9a"),(7.2,0.5,"Blackboard","#eee","#424242"),
                         (9.7,1.5,"ReportAgent","#e0f7fa","#00695c")]:
        _box(ax,(x,y),1.8,0.9,t,fc=fc,ec=ec,fs=8)
    _arr(ax,(2.1,1.95),(2.5,1.95)); _arr(ax,(4.3,2.1),(4.7,2.7)); _arr(ax,(4.3,1.8),(4.7,1.1))
    _arr(ax,(6.5,2.95),(7.2,2.95)); _arr(ax,(6.5,0.95),(7.2,0.95))
    _arr(ax,(9.0,2.7),(9.7,2.1)); _arr(ax,(9.0,1.15),(9.7,1.7))
    fig.savefig(p,bbox_inches="tight",facecolor="white"); plt.close(fig); return p

def gen_dual_engine():
    p=CHART_DIR/"fig_dual_engine.png"
    fig,ax=plt.subplots(figsize=(10,5.5),dpi=150); ax.set_xlim(0,10); ax.set_ylim(-0.8,5.5); ax.axis("off")
    ax.set_title("“财报SQL+研报RAG”双检索引擎", fontsize=13, fontweight="bold", pad=10)
    _box(ax,(3.8,4.2),2.4,0.8,"用户自然语言问题",fc="#e3f2fd",ec="#1565c0",fs=10)
    _box(ax,(3.8,3.0),2.4,0.7,"LLM 任务拆解",fc="#fff8e1",ec="#f57f17",fs=9)
    _box(ax,(0.5,1.5),2.2,0.8,"Text-to-SQL",fc="#e8f5e9",ec="#2e7d32")
    _box(ax,(0.5,0.2),2.2,0.8,"SQLite 双库验证",fc="#c8e6c9",ec="#2e7d32")
    _box(ax,(7.3,1.5),2.2,0.8,"TF-IDF 检索",fc="#f3e5f5",ec="#6a1b9a")
    _box(ax,(7.3,0.2),2.2,0.8,"研报片段召回",fc="#e1bee7",ec="#6a1b9a")
    _box(ax,(3.5,-0.5),3.0,0.7,"LLM 多源整合→回答",fc="#fff3e0",ec="#e65100",fs=9)
    _arr(ax,(5,4.2),(5,3.75)); _arr(ax,(3.8,3.2),(2.7,2.35)); _arr(ax,(6.2,3.2),(7.3,2.35))
    _arr(ax,(1.6,1.5),(1.6,1.05)); _arr(ax,(8.4,1.5),(8.4,1.05))
    _arr(ax,(1.6,0.2),(4.0,0.15)); _arr(ax,(8.4,0.2),(6.5,0.15))
    fig.savefig(p,bbox_inches="tight",facecolor="white"); plt.close(fig); return p

def gen_radar():
    p=CHART_DIR/"fig_radar.png"
    labels=["盈利能力","成长能力","运营效率","偿债能力","现金流质量"]
    angles=np.linspace(0,2*np.pi,5,endpoint=False).tolist()+[0]
    fig,ax=plt.subplots(figsize=(7,7),subplot_kw=dict(polar=True),dpi=150)
    for vals,c,lb in [([78.5,62.3,71,85.2,68.7],"#4e79a7","华润三九"),([45.2,38.1,55.6,60.3,42.8],"#e15759","金花股份")]:
        ax.plot(angles,vals+vals[:1],lw=2,label=lb,color=c); ax.fill(angles,vals+vals[:1],alpha=0.15,color=c)
    ax.set_thetagrids([a*180/np.pi for a in angles[:-1]],labels,fontsize=11)
    ax.set_ylim(0,100); ax.set_title("企业运营评估雷达图(2024年)",fontsize=14,fontweight="bold",pad=20)
    ax.legend(loc="upper right",bbox_to_anchor=(1.25,1.1)); ax.grid(True,alpha=0.3)
    fig.savefig(p,bbox_inches="tight",facecolor="white"); plt.close(fig); return p

def gen_test_acc():
    p=CHART_DIR/"fig_test_acc.png"
    cats=["SQL执行\n成功率","SQL结果\n正确率","KB Top-1\n相关率","KB Top-3\n相关率","双库一致率"]
    vals=[93.3,89.3,70.0,85.0,96.7]
    colors=["#4e79a7","#59a14f","#e15759","#f28e2b","#76b7b2"]
    fig,ax=plt.subplots(figsize=(9,5),dpi=150)
    bars=ax.bar(cats,vals,color=colors,width=0.55,edgecolor="white",lw=1.2)
    for b,v in zip(bars,vals): ax.text(b.get_x()+b.get_width()/2,b.get_height()+1.5,f"{v}%",ha="center",fontsize=11,fontweight="bold")
    ax.set_ylim(0,110); ax.set_ylabel("准确率(%)"); ax.set_title("核心功能测试准确率",fontsize=14,fontweight="bold")
    ax.spines["top"].set_visible(False); ax.spines["right"].set_visible(False); ax.grid(axis="y",alpha=0.3)
    fig.savefig(p,bbox_inches="tight",facecolor="white"); plt.close(fig); return p

def gen_resp_time():
    p=CHART_DIR/"fig_resp_time.png"
    funcs=["自然语言问数","运营评估","风险洞察","完整报告"]
    avg=[8.2,18.5,12.3,45.2]; ftk=[2.1,3.5,2.8,4.0]; x=np.arange(4); w=0.3
    fig,ax=plt.subplots(figsize=(9,5),dpi=150)
    b1=ax.bar(x-w/2,avg,w,label="平均响应时间",color="#4e79a7")
    b2=ax.bar(x+w/2,ftk,w,label="首字延迟",color="#f28e2b")
    for b in b1: ax.text(b.get_x()+b.get_width()/2,b.get_height()+0.8,f"{b.get_height()}s",ha="center",fontsize=9)
    for b in b2: ax.text(b.get_x()+b.get_width()/2,b.get_height()+0.8,f"{b.get_height()}s",ha="center",fontsize=9)
    ax.set_xticks(x); ax.set_xticklabels(funcs); ax.set_ylabel("时间(秒)")
    ax.set_title("各功能模块响应时间对比",fontsize=14,fontweight="bold"); ax.legend()
    ax.spines["top"].set_visible(False); ax.spines["right"].set_visible(False); ax.grid(axis="y",alpha=0.3)
    fig.savefig(p,bbox_inches="tight",facecolor="white"); plt.close(fig); return p

def gen_risk_engine():
    p=CHART_DIR/"fig_risk_engine.png"
    fig,ax=plt.subplots(figsize=(10,4.5),dpi=150); ax.set_xlim(0,10); ax.set_ylim(0,4.5); ax.axis("off")
    ax.set_title("规则引擎+LLM融合的风险洞察流程",fontsize=13,fontweight="bold",pad=10)
    _box(ax,(0.3,3.2),2.6,0.8,"第一层：量化规则\nSQL风险扫描",fc="#ffcdd2",ec="#c62828",fs=8)
    _box(ax,(3.6,3.2),2.8,0.8,"第二层：研报上下文\nTF-IDF检索",fc="#e1bee7",ec="#6a1b9a",fs=8)
    _box(ax,(7.1,3.2),2.6,0.8,"第三层：LLM研判\n角色化报告",fc="#c8e6c9",ec="#2e7d32",fs=8)
    _box(ax,(0.3,1.5),2.6,1.0,"风险信号\n业绩下滑·偿债·现金流",fc="#fff3e0",ec="#e65100",fs=7)
    _box(ax,(3.6,1.5),2.8,1.0,"机会信号\n研发突破·行业利好",fc="#e0f7fa",ec="#00695c",fs=7)
    _box(ax,(7.1,1.5),2.6,1.0,"输出：差异化\n风险/机会报告",fc="#f3e5f5",ec="#4a148c",fs=7)
    _arr(ax,(2.9,3.6),(3.6,3.6)); _arr(ax,(6.4,3.6),(7.1,3.6))
    _arr(ax,(1.6,3.2),(1.6,2.55)); _arr(ax,(5.0,3.2),(5.0,2.55)); _arr(ax,(8.4,3.2),(8.4,2.55))
    fig.savefig(p,bbox_inches="tight",facecolor="white"); plt.close(fig); return p

def gen_all_charts():
    return {
        "arch": gen_architecture(), "flow": gen_agent_flow(),
        "dual": gen_dual_engine(), "radar": gen_radar(),
        "test": gen_test_acc(), "resp": gen_resp_time(), "risk": gen_risk_engine(),
    }
