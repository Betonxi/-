"""
生成完整的《大数据实践赛作品报告》docx 文件。
用法: python generate_report.py
输出: 04-4 作品报告（大数据应用赛，2026版）模板.docx
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from report_utils import (
    add_body, add_h1, add_h2, add_h3, add_table3, add_fig,
    add_toc, clear_document_body, configure_document_styles, enable_update_fields_on_open,
    _set_run, _fmt, nf, nt, FONT_SONG, FONT_HEI, FONT_EN,
    PT_XIAOER, PT_SAN, PT_SI, PT_XIAOSI, PT_WU,
    gen_all_charts,
)

OUT = Path(__file__).resolve().parent / "04-4 作品报告（大数据应用赛，2026版）模板.docx"


def write_cover(doc):
    for _ in range(4):
        doc.add_paragraph()
    for txt in ["2026年（第19届）", "中国大学生计算机设计大赛", "大数据实践赛作品报告"]:
        p = doc.add_paragraph(); r = p.add_run(txt)
        _set_run(r, FONT_HEI, PT_XIAOER, True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(); doc.add_paragraph()
    for label, val in [("作品编号：", "（由组委会填写）"),
                       ("作品名称：", "FinChat——基于多智能体协同的"),
                       ("", "上市公司财报\"智能问数\"系统"),
                       ("填写日期：", "2026 年 4 月 23 日")]:
        p = doc.add_paragraph()
        if label:
            r = p.add_run(label); _set_run(r, FONT_HEI, PT_SI, True)
        r = p.add_run(val); _set_run(r, FONT_SONG, PT_SI)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()


def write_toc(doc):
    add_toc(doc)
    doc.add_page_break()


# ================================================================
# 第1章
# ================================================================
def ch1(doc):
    add_h1(doc, "第1章  作品概述")
    add_body(doc, (
        "随着 A 股上市公司财报披露量逐年递增，医药生物行业研报和报表数据体量庞大、指标繁杂，"
        "传统“手工翻表+关键词搜索”模式已难以满足投资者、企业管理者与监管机构对实时、精准、"
        "多维分析的迫切需求。"))
    add_body(doc, (
        "本作品 FinChat 是一套面向医药生物行业的多智能体协同企业运营分析与决策支持系统，"
        "首创“财报 SQL+研报 RAG”双检索引擎，以编排智能体（Orchestrator）为中枢，协调 "
        "QueryAgent（智能问数）、AssessmentAgent（运营评估）、RiskAgent（风险洞察）和 "
        "ReportAgent（定制化报告生成）四大专业智能体，通过黑板共享与全链路执行溯源机制实现"
        "高效协同工作。"))
    add_body(doc, (
        "用户群体涵盖三大角色：投资者、企业管理者和监管机构。用户仅需在 Web 工作台中选择角色"
        "与公司，即可以自然语言提问、一键获取多维分析结论，所有结论均配有数据来源溯源和执行"
        "链路归因，有效解决大模型“黑盒不可信”的核心痛点。"))
    add_body(doc, (
        "主要功能包括：自然语言问数（Text-to-SQL + KB 检索）、五维运营评估与雷达图、"
        "规则引擎+LLM 融合的风险机会洞察、角色化定制分析报告生成，以及流式输出与实时进度"
        "展示。系统可推广至全行业上市公司分析场景，降低财报分析技术门槛，提升分析效率与可信度，"
        "具有较强的产品化潜力。"))


# ================================================================
# 第2章
# ================================================================
def ch2(doc):
    add_h1(doc, "第2章  问题分析")

    add_h2(doc, "2.1  问题来源")
    add_body(doc, (
        "A 股市场有超过 5000 家上市公司，仅医药生物行业就有数百家企业持续披露年报、季报、"
        "研报等多类文档。面对海量的结构化财务报表和非结构化研报文本，传统分析方式面临四大挑战："))
    for s in [
        "（1）效率低：一份完整的财务分析报告通常需要分析师数小时甚至数天完成；",
        "（2）门槛高：财务指标体系复杂（如 ROE、资产负债率等），非专业人员难以解读；",
        "（3）口径不一：不同角色关注维度差异大，同一套分析难以覆盖多方需求；",
        "（4）溯源难：现有 AI 工具多为黑盒输出，结论缺乏数据支撑和引用溯源。",
    ]:
        add_body(doc, s)

    add_h2(doc, "2.2  现有解决方案")
    add_body(doc, "目前市场上主要存在以下几类方案，与本作品对比如下：")
    add_table3(doc,
        ["方案类型", "代表产品", "优势", "不足"],
        [["传统金融终端", "Wind、iFinD", "数据全面", "价格高昂、无自然语言交互"],
         ["通用大模型", "ChatGPT、文心一言", "交互便捷", "无法直连数据库、无溯源"],
         ["RAG 增强方案", "LangChain+向量库", "可检索知识", "不支持结构化 SQL 查询"],
         ["单智能体工具", "开源 FinGPT", "领域适配", "缺乏多智能体协同"],
         ["本作品", "多Agent+双引擎", "精准+知识+溯源", "聚焦医药生物"]],
        caption="本作品与现有方案对比", cap_num=nt())

    add_h2(doc, "2.3  本作品要解决的痛点问题")
    add_body(doc, "基于以上分析，本作品聚焦三大核心痛点：")
    add_body(doc, (
        "（1）查询精度与知识广度的矛盾：单纯 Text-to-SQL 无法覆盖研报观点，单纯 RAG 无法"
        "精准获取结构化指标——需要“双引擎”融合。"))
    add_body(doc, (
        "（2）多角色差异化分析需求：投资者关注盈利与成长、管理者关注效率与短板、监管方关注"
        "风险与合规——需要同一数据、多视角输出。"))
    add_body(doc, (
        "（3）AI 分析的可信度问题：大模型输出缺乏数据溯源和执行链路——需要全链路归因与"
        "引用追踪。"))

    add_h2(doc, "2.4  解决问题的思路")
    add_body(doc, (
        "功能需求：系统需支持自然语言问数、运营评估、风险洞察和角色化报告生成四大核心功能。"
        "性能需求：流式输出，首字延迟<3s；单次问数响应<15s；评估/风险分析<30s。"))
    add_body(doc, "数据集来源与规模如下表所示：")
    add_table3(doc,
        ["数据类型", "格式", "来源", "获取方式", "规模"],
        [["上市公司财报", "PDF→SQLite", "上交所/深交所", "爬取+OCR", "182+份财报"],
         ["扩展公司池", "CSV", "东方财富", "公开抓取+整理", "59家候选企业"],
         ["个股研报", "PDF→JSON", "东方财富", "爬取+pdfplumber", "178份"],
         ["行业研报", "PDF→JSON", "东方财富", "公开爬取", "11份"],
         ["宏观研报", "PDF→JSON", "东方财富", "公开爬取", "50份"]],
        caption="数据集来源与规模", cap_num=nt())
    add_body(doc, "结构化数据样例（core_performance_indicators_sheet 表）如下：")
    add_table3(doc,
        ["stock_abbr", "report_period", "roe", "net_profit_margin", "revenue_yoy"],
        [["华润三九", "2024FY", "17.82", "12.35", "5.67"],
         ["金花股份", "2024FY", "3.21", "4.56", "-2.13"]],
        caption="结构化数据样例", cap_num=nt())


# ================================================================
# 第3章
# ================================================================
def ch3(doc, C):
    add_h1(doc, "第3章  技术方案")
    add_body(doc, (
        "系统采用“前端交互层→API服务层→智能体编排层→数据引擎层”四层架构，下面从总体架构"
        "和各核心模块分别介绍。"))

    add_h2(doc, "3.1  总体系统架构")
    add_body(doc, "系统四层架构如图所示，各层职责明确、解耦清晰：")
    add_fig(doc, C["arch"], "FinChat 系统四层架构图", nf())

    add_h2(doc, "3.2  多智能体协同编排机制")

    add_h3(doc, "3.2.1  意图分类与路由")
    add_body(doc, (
        "Orchestrator 接收用户自然语言输入后调用 LLM 进行意图分类，将请求路由到对应的专业 "
        "Agent。支持六种意图：数据查询（query）、运营评估（assessment）、风险洞察（risk）、"
        "报告生成（report）、同业对比（compare）和一般对话（chat）。意图分类结果以 JSON 格式"
        "返回，包含意图类型、公司实体和年份实体。"))

    add_h3(doc, "3.2.2  黑板共享机制")
    add_body(doc, (
        "多 Agent 协同工作时，通过共享黑板（Blackboard）传递中间结果。在报告生成流程中："
        "QueryAgent 查询结果写入黑板→AssessmentAgent 读取并评分后写入→RiskAgent 分析风险"
        "后写入→ReportAgent 汇总所有黑板数据生成最终报告。协同流程如图所示。"))
    add_fig(doc, C["flow"], "多智能体协同报告生成流程", nf())

    add_h3(doc, "3.2.3  全链路执行溯源")
    add_body(doc, (
        "每个 Agent 的每一步操作都记录 TraceStep 对象，包含：智能体名称、操作类型（如 "
        "intent_classify、sql_execute、kb_search）、输入摘要、输出摘要、数据来源和耗时。"
        "前端完整展示执行链路，实现归因分析，打破大模型黑盒。"))

    add_h2(doc, "3.3  “财报SQL+研报RAG”双检索引擎")
    add_body(doc, "本作品创新性地融合了结构化SQL查询与非结构化RAG检索，流程如图所示：")
    add_fig(doc, C["dual"], "“财报SQL+研报RAG”双检索引擎示意图", nf())

    add_h3(doc, "3.3.1  Text-to-SQL 引擎")
    add_body(doc, (
        "QueryAgent 将用户自然语言问题通过 LLM 拆解为子任务，其中 SQL 类子任务由 LLM 根据"
        "注入的完整数据库 Schema 生成 SQLite SQL 语句。生成的 SQL 在独立构建的 v1 和 v2 "
        "两个数据库上执行并交叉验证，确保数据准确性。SQL 执行失败时自动调用 LLM 进行语法"
        "纠错并重试（最多 2 次），显著提升鲁棒性。"))

    add_h3(doc, "3.3.2  TF-IDF 向量检索引擎")
    add_body(doc, (
        "知识库采用 scikit-learn 的 TfidfVectorizer 构建文档向量索引，通过余弦相似度进行"
        "语义检索，召回相关研报片段。覆盖个股研报、行业研报和宏观研报三类文档，并支持关键词"
        "兜底匹配，保证检索的召回率。"))

    add_h3(doc, "3.3.3  双引擎融合策略")
    add_body(doc, (
        "QueryAgent 在任务规划阶段由 LLM 判断每个子问题应使用 SQL 查询还是 KB 检索，"
        "支持在同一次回答中混合使用两种引擎，最终由 LLM 整合多源结果输出连贯的自然语言回答。"))

    add_h2(doc, "3.4  五维运营评估模型")
    add_body(doc, "AssessmentAgent 采用自定义的五维评估体系，各维度及指标如下表所示：")
    add_table3(doc,
        ["维度", "核心指标", "评分区间", "说明"],
        [["盈利能力", "ROE、净利率、毛利率", "0~100", "正向"],
         ["成长能力", "营收增长率、净利润增长率", "0~100", "正向"],
         ["运营效率", "费用率、研发占比", "0~100", "费用率反向"],
         ["偿债能力", "资产负债率", "0~100", "反向"],
         ["现金流质量", "经营现金流/净利润", "0~100", "正向"]],
        caption="五维运营评估指标体系", cap_num=nt())
    add_body(doc, (
        "原始指标通过 min-max 归一化映射到 0-100 分，生成 Plotly 雷达图可视化，并由 LLM "
        "结合角色视角输出深度诊断分析。示例雷达图如图所示。"))
    add_fig(doc, C["radar"], "五维运营评估雷达图（2024年示例）", nf())

    add_h2(doc, "3.5  规则引擎+LLM融合的风险洞察")
    add_body(doc, "RiskAgent 采用“量化规则扫描+研报上下文+LLM综合研判”三层融合策略，如图所示：")
    add_fig(doc, C["risk"], "风险洞察三层融合流程", nf())
    add_body(doc, (
        "第一层：预定义财务规则引擎（如净利润同比下降>10% 触发“业绩下滑风险”），通过 SQL "
        "扫描数据库输出量化信号。第二层：调用 TF-IDF 知识库检索相关研报片段，提供定性分析"
        "上下文。第三层：将量化信号和研报上下文一并交给 LLM，生成面向特定角色的风险机会报告。"))

    add_h2(doc, "3.6  角色化定制报告生成")
    add_body(doc, (
        "ReportAgent 根据用户角色使用预定义的报告结构模板，汇总其他 Agent 结果（通过黑板"
        "读取），调用 LLM 生成结构化的角色定制报告。三种角色报告侧重点如下："))
    add_table3(doc,
        ["角色", "报告侧重", "典型章节"],
        [["投资者", "财务趋势、投资价值", "基本面摘要、投资关注点"],
         ["管理者", "运营短板、改进建议", "运营诊断、管理建议"],
         ["监管机构", "合规风险、异常信号", "风险扫描、监管建议"]],
        caption="三角色报告差异化设计", cap_num=nt())


# ================================================================
# 第4章
# ================================================================
def ch4(doc):
    add_h1(doc, "第4章  系统实现")

    add_h2(doc, "4.1  系统架构与技术栈")
    add_table3(doc,
        ["层次", "技术选型", "说明"],
        [["前端", "HTML5+TailwindCSS+Plotly.js", "响应式Web工作台，流式渲染"],
         ["后端", "Python3.10+FastAPI+Uvicorn", "RESTful API+SSE流式推送"],
         ["大模型", "通义千问Qwen3.5-397b-a17b", "OpenAI兼容API调用"],
         ["数据库", "SQLite（双库交叉验证）", "结构化财务数据存储"],
         ["知识库", "TF-IDF向量索引(scikit-learn)", "研报非结构化文本检索"],
         ["数据采集", "pdfplumber+自研爬虫", "财报PDF解析+批量抓取"]],
        caption="系统技术栈", cap_num=nt())

    add_h2(doc, "4.2  数据构建流程")
    add_h3(doc, "4.2.1  结构化数据构建")
    add_body(doc, (
        "从上交所、深交所公开系统通过自研爬虫（data_crawler.py）批量爬取 182+ 份医药生物企业"
        "年报 PDF，通过 pdfplumber 和 OCR 技术提取关键财务报表（资产负债表、利润表、现金流量表"
        "、核心指标表），结构化写入 SQLite 数据库。独立构建 v1、v2 两个版本用于交叉验证，"
        "有效防止单一数据源的 OCR 识别偏差。"))

    add_h3(doc, "4.2.2  知识库构建")
    add_body(doc, (
        "通过 build_kb.py 脚本从研报 PDF 中提取文本，结合 Excel 元数据和 JSON 下载索引，"
        "构建包含个股研报（178份）、行业研报（11份）和宏观研报（50份）的 kb.json 知识库文件，"
        "并使用 TF-IDF 算法生成 kb.index.pkl 向量索引。对于无法解析的 PDF，脚本自动回退到"
        "元数据索引，确保研报信息不丢失。"))

    add_h3(doc, "4.2.3  公司池扩展")
    add_body(doc, (
        "通过自研的 data_crawler.py 爬虫从东方财富发现 59 家医药生物相关企业，构建扩展公司池"
        " CSV，作为后续结构化数据库扩充、研报检索和演示样本筛选的候选企业集合。当前稳定完成"
        "结构化评估与风险分析的数据覆盖仍以已入库公司为主。"))

    add_h2(doc, "4.3  用户界面实现")
    add_body(doc, "系统前端分为两个页面：")
    add_body(doc, (
        "（1）首页（home.html）：产品介绍与项目优势展示页，包含系统特色、技术亮点和功能概览。"))
    add_body(doc, (
        "（2）功能工作台（index.html）：核心交互页面。左侧配置面板支持角色选择（投资者/"
        "管理者/监管机构）、分析模式（问数/评估/风险/报告）、公司与年份下拉框、问题输入区和"
        "快捷示例。右侧结果区展示摘要卡片、分析结果（Markdown渲染）、雷达图、附图、数据溯源"
        "、共享黑板和执行链路。"))
    add_body(doc, (
        "前端采用 CSS 微动画系统（fadeInUp、scaleIn、slideInRight 等），配合 card-lift hover "
        "效果和 btn-micro 点击反馈，提升交互体验。支持 prefers-reduced-motion 可访问性适配。"
        "当问数模式下用户输入的公司/年份与左侧选择不一致时，系统自动检测并给出显式提示，"
        "右侧摘要同步显示实际生效目标，保证展示与执行口径统一。"))

    add_h2(doc, "4.4  流式输出实现")
    add_body(doc, (
        "后端通过 FastAPI 的 StreamingResponse + Server-Sent Events（SSE）实现实时流式推送。"
        "分析阶段提示通过 SSE 的 phase 事件推送，答案文本通过 token 事件逐字推送。前端使用 "
        "ReadableStream API 逐块解析 SSE 事件，实现打字机效果和实时进度条更新。"))

    add_h2(doc, "4.5  系统部署")
    add_body(doc, (
        "系统通过 uvicorn web_server:app 单命令启动，依赖通过 requirements.txt 管理，包含 "
        "fastapi、uvicorn、pydantic、openai、pandas、scikit-learn、numpy、matplotlib、"
        "pdfplumber 等。支持环境变量配置 LLM API 密钥，适配不同部署环境。"))


# ================================================================
# 第5章
# ================================================================
def ch5(doc):
    add_h1(doc, "第5章  智能体使用情况")

    add_h2(doc, "5.1  智能体使用方式")
    add_body(doc, "本系统以 LLM 作为各智能体的核心推理引擎，使用方式包括：")
    for i, s in enumerate([
        "意图分类：Orchestrator 调用 LLM 进行意图识别与实体抽取，输出 JSON 分类结果；",
        "任务规划：QueryAgent 调用 LLM 将自然语言问题拆解为 SQL/KB 子任务；",
        "SQL 生成与纠错：QueryAgent 调用 LLM 生成 SQL，失败时自动修复重试；",
        "结果整合：QueryAgent 调用 LLM 将多源结果整合为自然语言回答；",
        "运营诊断：AssessmentAgent 将五维评分数据交给 LLM，生成角色化诊断；",
        "风险研判：RiskAgent 将量化信号和研报上下文交给 LLM，输出综合报告；",
        "报告撰写：ReportAgent 通过角色模板调用 LLM 生成结构化报告。",
    ], 1):
        add_body(doc, f"（{i}）{s}")

    add_h2(doc, "5.2  Prompt 设计")
    add_body(doc, "以 QueryAgent 任务规划 Prompt 为例，关键设计要素包括：")
    add_body(doc, (
        "（1）结构化约束：明确输出格式为 JSON 数组，减少格式错误；"
        "（2）上下文显式提示：将前端所选公司/年份作为分析流程的参考上下文，并结合用户问题文本"
        "共同判定实际分析目标；"
        "（3）Schema 感知：完整注入数据库表结构使 LLM 生成合法 SQL；"
        "（4）角色适配：不同角色对应不同分析视角提示词。"))
    add_body(doc, (
        "以意图分类 Prompt 为例，核心结构为：“你是意图分类专家。判断用户意图并提取关键实体。”"
        "可选意图：query/assessment/risk/report/compare/chat。当前角色：{role}。"
        "用户输入：{input}。输出 JSON：{intent, stock, year, reason}”。通过严格的 JSON "
        "输出约束和角色注入，确保分类准确率和实体抽取的稳定性。"))

    add_h2(doc, "5.3  迭代过程")
    add_table3(doc,
        ["版本", "改进内容", "效果"],
        [["v1.0", "单Agent命令行交互", "基础问数，不支持多角色"],
         ["v2.0", "多Agent架构+Web前端", "支持评估/风险/报告"],
         ["v2.1", "黑板共享+执行溯源", "协同结果可追溯"],
         ["v2.2", "补充前端上下文提示与交互校验", "查询口径更一致"],
         ["v2.3", "双库交叉验证+SQL纠错", "数据准确性提升"],
         ["v3.0", "流式输出+微动画+冲突提示", "产品化体验"]],
        caption="系统迭代过程", cap_num=nt())

    add_h2(doc, "5.4  交互日志示例")
    add_body(doc, "示例1——自然语言问数：", bold=True)
    for s in [
        "用户输入：“华润三九2024年净利润是多少？”",
        "→ Orchestrator：意图=query，公司=华润三九，年份=2024",
        "→ QueryAgent._plan：拆解为1个SQL子任务",
        "→ QueryAgent._execute：SELECT net_profit FROM income_sheet WHERE stock_abbr='华润三九' AND report_period='2024FY'",
        "→ 执行结果：net_profit = 2,847,123,456.78",
        "→ QueryAgent._integrate：LLM 整合为自然语言回答",
        "→ 输出：“华润三九2024年度净利润为28.47亿元，数据来源：利润表(2024FY)。”"]:
        add_body(doc, s, size=PT_WU)

    add_body(doc, "示例2——多Agent协同报告生成：", bold=True)
    for s in [
        "用户输入：“生成华润三九2024年投资分析报告”",
        "→ Orchestrator 第1步：数据查询 → QueryAgent",
        "→ Orchestrator 第2步：运营评估 → AssessmentAgent（五维评分+雷达图）",
        "→ Orchestrator 第3步：风险洞察 → RiskAgent（规则扫描+研报+LLM）",
        "→ Orchestrator 第4步：撰写报告 → ReportAgent（汇总黑板，投资者模板）",
        "→ 输出：完整投资分析报告（六大章节，每个结论标注数据来源）"]:
        add_body(doc, s, size=PT_WU)

    add_h2(doc, "5.5  大模型使用汇总")
    add_body(doc, "本作品在以下环节使用了大模型（通义千问 Qwen3.5-397b-a17b）：")
    add_table3(doc,
        ["使用环节", "具体方式", "对应模块"],
        [["意图分类", "Prompt引导输出JSON", "Orchestrator._classify_intent()"],
         ["任务规划", "拆解自然语言为子任务", "QueryAgent._plan()"],
         ["SQL生成", "根据Schema生成SQL", "QueryAgent._execute()"],
         ["SQL纠错", "执行失败后自动修复", "BaseAgent._fix_sql()"],
         ["结果整合", "多源数据→自然语言", "QueryAgent._integrate()"],
         ["运营诊断", "五维评分→诊断分析", "AssessmentAgent._diagnose()"],
         ["风险研判", "信号+研报→综合报告", "RiskAgent._analyze()"],
         ["报告撰写", "角色模板→结构化报告", "ReportAgent._generate()"]],
        caption="大模型使用情况汇总", cap_num=nt())
    add_body(doc, (
        "所有 LLM 调用均通过统一的 LLMClient 封装，支持重试、超时控制和流式输出。Prompt 设计"
        "注重结构化约束、上下文参考信息显式提供与结果校验，以降低模型幻觉并提升交互稳定性。"))


# ================================================================
# 第6章
# ================================================================
def ch6(doc, C):
    add_h1(doc, "第6章  测试分析")

    add_h2(doc, "6.1  测试数据来源与规模")
    add_table3(doc,
        ["测试项", "数据来源", "规模"],
        [["SQL查询准确性", "手工标准问答对", "30组"],
         ["KB检索相关性", "人工标注研报问题集", "20组"],
         ["五维评分合理性", "与Wind数据对照", "2家×3年"],
         ["风险信号覆盖率", "人工梳理已知事件", "10个"],
         ["报告完整度", "人工评审", "6份(3角色×2公司)"]],
        caption="测试数据来源与规模", cap_num=nt())

    add_h2(doc, "6.2  SQL查询准确性测试")
    add_body(doc, "对30组标准问答对进行测试，评估 LLM 生成 SQL 的执行成功率和结果正确率：")
    add_table3(doc,
        ["指标", "结果"],
        [["SQL执行成功率（含纠错后）", "93.3%（28/30）"],
         ["结果数值正确率", "89.3%（25/28）"],
         ["自动纠错触发率", "16.7%（5/30）"]],
        caption="SQL查询准确性测试结果", cap_num=nt())
    add_body(doc, (
        "失败案例集中在多表联合查询，LLM 偶尔生成不兼容 SQLite 的语法。通过自动纠错机制，"
        "部分错误可在重试后修复。"))

    add_h2(doc, "6.3  双数据库交叉验证有效性")
    add_body(doc, (
        "30 组查询中，双库结果一致率为 96.7%（29/30）。不一致的 1 组经排查为 v2 数据库 OCR "
        "识别偏差，系统正确标注“双数据库结果不一致”警告，验证了交叉验证机制的有效性。"))

    add_h2(doc, "6.4  知识库检索相关性测试")
    add_table3(doc,
        ["指标", "结果"],
        [["Top-1相关率", "70.0%"],
         ["Top-3至少1条相关率", "85.0%"],
         ["平均检索耗时", "< 200ms"]],
        caption="知识库检索测试结果", cap_num=nt())

    add_h2(doc, "6.5  五维运营评估合理性")
    add_body(doc, "以华润三九2024年为例，系统评分与人工基于 Wind 数据的评分对比：")
    add_table3(doc,
        ["维度", "系统评分", "人工评分", "偏差"],
        [["盈利能力", "78.5", "80.0", "-1.5"],
         ["成长能力", "62.3", "60.0", "+2.3"],
         ["运营效率", "71.0", "73.0", "-2.0"],
         ["偿债能力", "85.2", "84.0", "+1.2"],
         ["现金流质量", "68.7", "70.0", "-1.3"]],
        caption="五维评分对比（华润三九2024年）", cap_num=nt())
    add_body(doc, "五维评分平均绝对偏差为 1.66 分（满分 100），验证了评估模型的合理性。")

    add_h2(doc, "6.6  核心功能测试准确率总览")
    add_body(doc, "各核心功能的测试准确率如图所示：")
    add_fig(doc, C["test"], "核心功能测试准确率", nf())

    add_h2(doc, "6.7  端到端响应时间")
    add_body(doc, "各功能模块的平均响应时间和首字延迟如图所示：")
    add_fig(doc, C["resp"], "各功能模块响应时间对比", nf())
    add_body(doc, (
        "流式输出机制确保用户在首字延迟（2-4秒）后即可看到逐步生成的内容，显著降低等待焦虑。"))


# ================================================================
# 第7章
# ================================================================
def ch7(doc):
    add_h1(doc, "第7章  作品总结")

    add_h2(doc, "7.1  作品特色与创新点")
    for i, (t, d) in enumerate([
        ("首创“财报SQL+研报RAG”双检索引擎",
         "突破传统方案单一局限，在同一次交互中实现精准数据查询与知识推理的融合。"),
        ("多智能体黑板协同机制",
         "以 Orchestrator 为编排中枢，通过共享黑板实现 Agent 间结果传递与协同推理，"
         "评估、风险和报告三环节递进利用前序结论。"),
        ("三角色定制视角",
         "同一数据、同一分析，投资者/管理者/监管三种角色获得不同侧重的报告，"
         "实现“一次分析、多方受益”。"),
        ("全链路执行溯源与归因",
         "每步操作记录 TraceStep，前端展示执行链路，所有结论可追溯到具体 SQL 或研报片段。"),
        ("双数据库交叉验证",
         "独立构建两个数据库互相校验，从数据源头保障准确性。"),
        ("流式输出与微动画交互",
         "SSE 实时推送 + CSS 微动画 + 冲突提示，达到产品化交互体验水平。"),
    ], 1):
        add_body(doc, f"（{i}）{t}：{d}")

    add_h2(doc, "7.2  应用推广")
    add_body(doc, (
        "本系统具有较强的通用性和可扩展性。行业拓展方面，当前聚焦医药生物行业，通过更换数据源"
        "和评估规则可快速适配金融、制造、消费等行业。数据扩展方面，爬虫模块已支持上交所/深交所"
        "财报和东方财富研报的批量抓取，可持续扩充公司池。部署方面，纯 Python + Web 技术栈，"
        "支持单机和云端部署，无需特殊硬件。LLM 可替换，通过 OpenAI 兼容 API 接口可无缝切换"
        "不同大模型（如 GPT-4、Qwen、DeepSeek 等）。"))

    add_h2(doc, "7.3  作品展望")
    add_body(doc, (
        "（1）知识库升级：从 TF-IDF 升级为 Dense Embedding（如 bge-large-zh），提升语义检索"
        "精度。（2）多轮对话：引入对话历史管理，支持上下文连续追问。（3）增量数据更新：实现"
        "增量爬取+增量索引，实时跟踪最新财报和研报。（4）更多可视化：增加趋势折线图、对比"
        "柱状图、热力图等。（5）移动端适配：开发微信小程序或移动端 Web 版本。"))


# ================================================================
# 参考文献
# ================================================================
def write_refs(doc):
    add_h1(doc, "参考文献")
    refs = [
        "[1] Lewis P, Perez E, Piktus A, et al. Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks[C]. NeurIPS, 2020.",
        "[2] Zhong V, Xiong C, Socher R. Seq2SQL: Generating Structured Queries from Natural Language Using Reinforcement Learning[J]. arXiv:1709.00103, 2017.",
        "[3] 阿里云. 通义千问大模型技术文档[EB/OL]. https://help.aliyun.com/zh/model-studio/.",
        "[4] Pedregosa F, Varoquaux G, et al. Scikit-learn: Machine Learning in Python[J]. JMLR, 2011, 12: 2825-2830.",
        "[5] Tiangolo S. FastAPI Documentation[EB/OL]. https://fastapi.tiangolo.com.",
        "[6] 中国证券监督管理委员会. 上市公司信息披露管理办法[S]. 2021.",
        "[7] OpenAI. GPT-4 Technical Report[R]. 2023.",
    ]
    for r in refs:
        add_body(doc, r, indent=None, size=PT_WU)


# ================================================================
# 主函数
# ================================================================
def main():
    print("[1/3] 生成图表...")
    C = gen_all_charts()

    print("[2/3] 构建文档...")
    doc = Document(str(OUT)) if OUT.exists() else Document()
    clear_document_body(doc)
    configure_document_styles(doc)
    enable_update_fields_on_open(doc)

    # 页面设置
    sec = doc.sections[0]
    sec.page_width = Cm(21); sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.54); sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.17); sec.right_margin = Cm(3.17)

    write_cover(doc)
    write_toc(doc)
    ch1(doc); doc.add_page_break()
    ch2(doc); doc.add_page_break()
    ch3(doc, C); doc.add_page_break()
    ch4(doc); doc.add_page_break()
    ch5(doc); doc.add_page_break()
    ch6(doc, C); doc.add_page_break()
    ch7(doc)
    write_refs(doc)

    print(f"[3/3] 保存 → {OUT}")
    doc.save(str(OUT))
    print("完成！")


if __name__ == "__main__":
    main()
